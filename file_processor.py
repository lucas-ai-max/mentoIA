"""
Módulo para processar e extrair texto de diferentes formatos de arquivo
"""
from typing import Optional
import io

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extrai texto de um arquivo baseado em sua extensão
    
    Args:
        file_content: Conteúdo binário do arquivo
        filename: Nome do arquivo (para determinar o tipo)
        
    Returns:
        Texto extraído do arquivo
        
    Raises:
        ValueError: Se o formato do arquivo não for suportado
    """
    filename_lower = filename.lower()
    
    # Arquivo de texto
    if filename_lower.endswith('.txt'):
        try:
            # Tentar diferentes encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            # Se nenhum encoding funcionou, usar utf-8 com erros ignorados
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            raise ValueError(f"Erro ao ler arquivo TXT: {str(e)}")
    
    # PDF
    elif filename_lower.endswith('.pdf'):
        try:
            import pypdf
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            if not text.strip():
                raise ValueError("Não foi possível extrair texto do PDF. O arquivo pode estar protegido ou ser uma imagem.")
            return text
        except ImportError:
            raise ValueError("Biblioteca pypdf não está instalada. Execute: pip install pypdf")
        except Exception as e:
            raise ValueError(f"Erro ao ler arquivo PDF: {str(e)}")
    
    # DOCX (Word)
    elif filename_lower.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Também extrair texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            if not text.strip():
                raise ValueError("Não foi possível extrair texto do arquivo DOCX")
            return text
        except ImportError:
            raise ValueError("Biblioteca python-docx não está instalada. Execute: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Erro ao ler arquivo DOCX: {str(e)}")
    
    # DOC (Word antigo) - tentar ler como DOCX primeiro
    elif filename_lower.endswith('.doc'):
        raise ValueError("Arquivos .doc (Word antigo) não são suportados diretamente. Por favor, converta para .docx ou .txt primeiro.")
    
    else:
        # Tentar como texto genérico
        try:
            return file_content.decode('utf-8', errors='ignore')
        except:
            raise ValueError(f"Formato de arquivo não suportado: {filename}. Formatos suportados: .txt, .pdf, .docx")

